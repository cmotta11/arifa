"""BVI Memorandum of Association builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("bvi_memorandum")
class MemorandumBuilder(BaseDocumentBuilder):
    """Build a BVI Memorandum of Association.

    Required context fields:
        entity_name, incorporation_date, registered_agent, registered_office.
    Optional:
        objects (str), authorized_shares (int or str).
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "incorporation_date",
        "registered_agent",
        "registered_office",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "memorandum.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "incorporation_date": context["incorporation_date"],
            "registered_agent": context["registered_agent"],
            "registered_office": context["registered_office"],
            "objects": context.get("objects", ""),
            "authorized_shares": context.get("authorized_shares", ""),
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Title
        title = doc.add_heading("MEMORANDUM OF ASSOCIATION", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(
            "BVI Business Companies Act, 2004 (as amended)"
        )
        run.italic = True

        entity_p = doc.add_paragraph()
        entity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = entity_p.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Clause 1 - Name
        doc.add_heading("1. Name", level=1)
        doc.add_paragraph(
            f"The name of the Company is {variables['entity_name']}."
        )

        # Clause 2 - Registered Office and Agent
        doc.add_heading("2. Registered Office and Agent", level=1)
        doc.add_paragraph(
            f"The registered office of the Company is at "
            f"{variables['registered_office']}."
        )
        doc.add_paragraph(
            f"The registered agent of the Company is "
            f"{variables['registered_agent']}."
        )

        # Clause 3 - Objects
        doc.add_heading("3. Objects", level=1)
        objects_text = (
            variables["objects"]
            or "The objects for which the Company is established are "
            "unrestricted and the Company shall have full power and "
            "authority to carry out any object not prohibited by the "
            "BVI Business Companies Act, 2004."
        )
        doc.add_paragraph(objects_text)

        # Clause 4 - Authorized Shares
        doc.add_heading("4. Authorised Shares", level=1)
        if variables["authorized_shares"]:
            doc.add_paragraph(
                f"The Company is authorised to issue a maximum of "
                f"{variables['authorized_shares']} shares."
            )
        else:
            doc.add_paragraph(
                "The Company is authorised to issue shares as determined "
                "by its Articles of Association."
            )

        # Clause 5 - Date
        doc.add_heading("5. Incorporation Date", level=1)
        doc.add_paragraph(
            f"Date of Incorporation: {variables['incorporation_date']}."
        )

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph("Signed by the Subscriber(s):")
        doc.add_paragraph()

        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig.add_run("_" * 40)
        doc.add_paragraph("Authorised Signatory")

        return self._save_docx_to_bytes(doc)
