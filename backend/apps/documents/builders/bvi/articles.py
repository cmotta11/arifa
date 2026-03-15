"""BVI Articles of Association builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("bvi_articles")
class ArticlesBuilder(BaseDocumentBuilder):
    """Build BVI Articles of Association.

    Required context fields:
        entity_name.
    Optional:
        share_classes (list), directors (list), meetings_config (dict),
        powers (str or list).
    """

    REQUIRED_FIELDS = ["entity_name"]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "articles.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        powers = context.get("powers", "")
        if isinstance(powers, list):
            powers = "; ".join(powers)

        return {
            "entity_name": context["entity_name"],
            "share_classes": context.get("share_classes") or [],
            "directors": context.get("directors") or [],
            "meetings_config": context.get("meetings_config") or {},
            "powers": powers,
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Title
        title = doc.add_heading("ARTICLES OF ASSOCIATION", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Part 1 - Interpretation
        doc.add_heading("PART 1 - INTERPRETATION", level=1)
        doc.add_paragraph(
            "In these Articles, unless the context otherwise requires, "
            "words and expressions defined in the Memorandum of Association "
            "and the BVI Business Companies Act, 2004 have the same meanings."
        )

        # Part 2 - Shares
        doc.add_heading("PART 2 - SHARES", level=1)

        share_classes = variables["share_classes"]
        if share_classes:
            doc.add_paragraph("The Company may issue the following classes of shares:")
            for sc in share_classes:
                if isinstance(sc, dict):
                    doc.add_paragraph(
                        f"  - Class: {sc.get('name', 'N/A')}, "
                        f"Shares: {sc.get('num_shares', 'N/A')}, "
                        f"Par Value: {sc.get('par_value', 'No par value')}",
                    )
                else:
                    doc.add_paragraph(f"  - {sc}")
        else:
            doc.add_paragraph(
                "The Company is authorised to issue shares of a single class "
                "with no par value."
            )

        # Part 3 - Directors
        doc.add_heading("PART 3 - DIRECTORS", level=1)

        directors = variables["directors"]
        if directors:
            doc.add_paragraph("The first directors of the Company shall be:")
            for director in directors:
                if isinstance(director, dict):
                    doc.add_paragraph(
                        f"  - {director.get('name', 'N/A')}"
                        f" ({director.get('position', 'Director')})"
                    )
                else:
                    doc.add_paragraph(f"  - {director}")
        else:
            doc.add_paragraph(
                "The Company shall have at least one director appointed "
                "in accordance with these Articles."
            )

        # Part 4 - Meetings
        doc.add_heading("PART 4 - MEETINGS OF MEMBERS", level=1)
        meetings_config = variables["meetings_config"]
        quorum = meetings_config.get("quorum", "a majority of the voting shares")
        notice_days = meetings_config.get("notice_days", 10)
        doc.add_paragraph(
            f"A quorum for a meeting of members consists of {quorum}."
        )
        doc.add_paragraph(
            f"Notice of at least {notice_days} days shall be given for "
            f"any meeting of members."
        )

        # Part 5 - Powers
        doc.add_heading("PART 5 - POWERS OF THE COMPANY", level=1)
        powers_text = (
            variables["powers"]
            or "The Company has, irrespective of corporate benefit, full "
            "capacity to carry on or undertake any business or activity, "
            "do any act or enter into any transaction."
        )
        doc.add_paragraph(powers_text)

        return self._save_docx_to_bytes(doc)
