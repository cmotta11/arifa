"""Panama Protocolo (Notary Protocol) builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_protocolo")
class ProtocoloBuilder(BaseDocumentBuilder):
    """Build a Panama notary protocol document.

    Required context fields:
        entity_name, deed_number, notary_name, date.
    Optional:
        witnesses (list), protocol_text (str).
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "deed_number",
        "notary_name",
        "date",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "protocolo.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "deed_number": context["deed_number"],
            "notary_name": context["notary_name"],
            "date": context["date"],
            "witnesses": context.get("witnesses") or [],
            "protocol_text": context.get("protocol_text", ""),
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Header
        heading = doc.add_heading("ESCRITURA PUBLICA", level=0)
        heading.alignment = WD_ALIGN_PARAGRAPH.CENTER

        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run(f"Numero de Escritura: {variables['deed_number']}")
        run.bold = True
        run.font.size = Pt(12)

        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"En la ciudad de Panama, a los {variables['date']}, ante mi, "
            f"{variables['notary_name']}, Notario Publico del Circuito de Panama, "
            f"comparece la sociedad {variables['entity_name']}."
        )

        if variables["protocol_text"]:
            doc.add_paragraph()
            doc.add_paragraph(variables["protocol_text"])

        # Witnesses
        if variables["witnesses"]:
            doc.add_paragraph()
            doc.add_heading("Testigos", level=2)
            for witness in variables["witnesses"]:
                if isinstance(witness, dict):
                    doc.add_paragraph(
                        f"  - {witness.get('name', 'N/A')}, "
                        f"Cedula: {witness.get('id_number', 'N/A')}"
                    )
                else:
                    doc.add_paragraph(f"  - {witness}")

        # Signature block
        doc.add_paragraph()
        doc.add_paragraph()
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig.add_run("_" * 40)
        notary_line = doc.add_paragraph()
        notary_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        notary_line.add_run(f"{variables['notary_name']}")
        title_line = doc.add_paragraph()
        title_line.alignment = WD_ALIGN_PARAGRAPH.CENTER
        title_line.add_run("Notario Publico")

        return self._save_docx_to_bytes(doc)
