"""Panama Director Resignation Letter builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_director_resignation")
class DirectorResignationBuilder(BaseDocumentBuilder):
    """Build a Panama director resignation letter.

    Required context fields:
        entity_name, director_name, resignation_date, effective_date.
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "director_name",
        "resignation_date",
        "effective_date",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "director_resignation.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "director_name": context["director_name"],
            "resignation_date": context["resignation_date"],
            "effective_date": context["effective_date"],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Date
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(variables["resignation_date"])

        doc.add_paragraph()

        # Addressee
        doc.add_paragraph("A la Junta Directiva de")
        entity_p = doc.add_paragraph()
        run = entity_p.add_run(variables["entity_name"])
        run.bold = True

        doc.add_paragraph()

        # Title
        title = doc.add_paragraph()
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = title.add_run("CARTA DE RENUNCIA")
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"Por medio de la presente, yo, {variables['director_name']}, "
            f"presento mi renuncia irrevocable como Director de "
            f"{variables['entity_name']}."
        )

        doc.add_paragraph(
            f"Esta renuncia sera efectiva a partir del "
            f"{variables['effective_date']}."
        )

        doc.add_paragraph()
        doc.add_paragraph(
            "Declaro que no tengo reclamacion alguna contra la sociedad "
            "por concepto de honorarios, salarios, o cualquier otra "
            "compensacion derivada de mi cargo como Director."
        )

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()
        doc.add_paragraph()

        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig.add_run("_" * 40)
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = name_p.add_run(variables["director_name"])
        run.bold = True
        label_p = doc.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_p.add_run("Director Renunciante")

        return self._save_docx_to_bytes(doc)
