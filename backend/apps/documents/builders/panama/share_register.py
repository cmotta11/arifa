"""Panama Share Register builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_share_register")
class ShareRegisterBuilder(BaseDocumentBuilder):
    """Build a Panama share register.

    Required context fields:
        entity_name.
    Optional:
        share_classes (list of dicts with ``name`` and ``issuances`` list).
    """

    REQUIRED_FIELDS = ["entity_name"]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "share_register.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "share_classes": context.get("share_classes") or [],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(10)
        style.font.name = "Times New Roman"

        title = doc.add_heading("REGISTRO DE ACCIONES", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        entity_p = doc.add_paragraph()
        entity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = entity_p.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        share_classes = variables["share_classes"]
        if not share_classes:
            doc.add_paragraph("No share classes have been registered.")
            return self._save_docx_to_bytes(doc)

        for sc in share_classes:
            class_name = sc.get("name", "N/A") if isinstance(sc, dict) else str(sc)
            issuances = sc.get("issuances", []) if isinstance(sc, dict) else []

            doc.add_heading(f"Clase: {class_name}", level=2)

            if not issuances:
                doc.add_paragraph("No issuances recorded for this class.")
                continue

            # Create table
            headers = [
                "Cert. No.",
                "Shareholder",
                "Shares",
                "Date Issued",
                "Date Cancelled",
            ]
            table = doc.add_table(rows=1, cols=len(headers))
            table.style = "Table Grid"

            # Header row
            for idx, header in enumerate(headers):
                cell = table.rows[0].cells[idx]
                cell.text = header
                for paragraph in cell.paragraphs:
                    for r in paragraph.runs:
                        r.bold = True

            # Data rows
            for issuance in issuances:
                row = table.add_row()
                if isinstance(issuance, dict):
                    row.cells[0].text = str(issuance.get("certificate_number", ""))
                    row.cells[1].text = str(issuance.get("shareholder", ""))
                    row.cells[2].text = str(issuance.get("num_shares", ""))
                    row.cells[3].text = str(issuance.get("date_issued", ""))
                    row.cells[4].text = str(issuance.get("date_cancelled", ""))

            doc.add_paragraph()

        return self._save_docx_to_bytes(doc)
