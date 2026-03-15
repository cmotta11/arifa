"""Panama Pacto Social (Articles of Incorporation) builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_pacto_social")
class PactoSocialBuilder(BaseDocumentBuilder):
    """Build a Panama corporation Pacto Social (Articles of Incorporation).

    Required context fields:
        entity_name, incorporation_date, jurisdiction, authorized_capital,
        capital_currency, share_classes (list), directors (list),
        resident_agent, registered_office, purposes (str or list).
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "incorporation_date",
        "jurisdiction",
        "authorized_capital",
        "capital_currency",
        "directors",
        "resident_agent",
        "registered_office",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "pacto_social.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        share_classes = context.get("share_classes") or []
        directors = context.get("directors") or []
        purposes = context.get("purposes", "")
        if isinstance(purposes, list):
            purposes = "; ".join(purposes)

        return {
            "entity_name": context["entity_name"],
            "incorporation_date": context["incorporation_date"],
            "jurisdiction": context["jurisdiction"],
            "authorized_capital": context["authorized_capital"],
            "capital_currency": context["capital_currency"],
            "share_classes": share_classes,
            "directors": directors,
            "resident_agent": context["resident_agent"],
            "registered_office": context["registered_office"],
            "purposes": purposes,
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            import io

            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        # Programmatic generation when no template file exists
        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Title
        title = doc.add_heading("PACTO SOCIAL", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        subtitle = doc.add_paragraph()
        subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = subtitle.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(14)

        doc.add_paragraph()

        # Article I - Name and Domicile
        doc.add_heading("ARTICULO PRIMERO - Nombre y Domicilio", level=1)
        doc.add_paragraph(
            f"La sociedad se denomina {variables['entity_name']}, constituida "
            f"de acuerdo con las leyes de la {variables['jurisdiction']}."
        )
        doc.add_paragraph(
            f"Fecha de constitucion: {variables['incorporation_date']}."
        )
        doc.add_paragraph(
            f"Oficina registrada: {variables['registered_office']}."
        )

        # Article II - Purposes
        doc.add_heading("ARTICULO SEGUNDO - Objeto Social", level=1)
        purposes_text = variables["purposes"] or "General commercial activities."
        doc.add_paragraph(purposes_text)

        # Article III - Capital
        doc.add_heading("ARTICULO TERCERO - Capital Social", level=1)
        doc.add_paragraph(
            f"El capital social autorizado es de {variables['capital_currency']} "
            f"{variables['authorized_capital']}."
        )
        if variables["share_classes"]:
            for sc in variables["share_classes"]:
                if isinstance(sc, dict):
                    doc.add_paragraph(
                        f"  - Clase: {sc.get('name', 'N/A')}, "
                        f"Acciones: {sc.get('num_shares', 'N/A')}, "
                        f"Valor nominal: {sc.get('par_value', 'N/A')}",
                    )
                else:
                    doc.add_paragraph(f"  - {sc}")

        # Article IV - Directors
        doc.add_heading("ARTICULO CUARTO - Junta Directiva", level=1)
        for director in variables["directors"]:
            if isinstance(director, dict):
                doc.add_paragraph(
                    f"  - {director.get('name', 'N/A')}"
                    f" ({director.get('position', 'Director')})"
                )
            else:
                doc.add_paragraph(f"  - {director}")

        # Article V - Resident Agent
        doc.add_heading("ARTICULO QUINTO - Agente Residente", level=1)
        doc.add_paragraph(
            f"El agente residente de la sociedad es {variables['resident_agent']}."
        )

        return self._save_docx_to_bytes(doc)
