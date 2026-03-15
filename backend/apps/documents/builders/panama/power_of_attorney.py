"""Panama Power of Attorney builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_power_of_attorney")
class PowerOfAttorneyBuilder(BaseDocumentBuilder):
    """Build a Panama power of attorney document.

    Required context fields:
        grantor_name, grantee_name, scope, date.
    Optional:
        witnesses (list).
    """

    REQUIRED_FIELDS = [
        "grantor_name",
        "grantee_name",
        "scope",
        "date",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "power_of_attorney.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "grantor_name": context["grantor_name"],
            "grantee_name": context["grantee_name"],
            "scope": context["scope"],
            "date": context["date"],
            "witnesses": context.get("witnesses") or [],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        title = doc.add_heading("PODER ESPECIAL", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"En la ciudad de Panama, a los {variables['date']}, "
            f"yo, {variables['grantor_name']}, por medio del presente documento, "
            f"confiero PODER ESPECIAL amplio y suficiente a "
            f"{variables['grantee_name']} para que en mi nombre y "
            f"representacion pueda:"
        )

        doc.add_paragraph()
        doc.add_paragraph(variables["scope"])

        doc.add_paragraph()
        doc.add_paragraph(
            "Este poder es valido desde la fecha de su otorgamiento y "
            "permanecera vigente hasta que sea revocado expresamente por "
            "el otorgante."
        )

        # Signatures
        doc.add_paragraph()
        doc.add_paragraph()

        # Grantor
        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig.add_run("_" * 40)
        name_p = doc.add_paragraph()
        name_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name_p.add_run(variables["grantor_name"])
        label_p = doc.add_paragraph()
        label_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        label_p.add_run("Otorgante")

        # Witnesses
        if variables["witnesses"]:
            doc.add_paragraph()
            doc.add_heading("Testigos:", level=2)
            for witness in variables["witnesses"]:
                doc.add_paragraph()
                w_sig = doc.add_paragraph()
                w_sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
                w_sig.add_run("_" * 40)
                w_name = doc.add_paragraph()
                w_name.alignment = WD_ALIGN_PARAGRAPH.CENTER
                if isinstance(witness, dict):
                    w_name.add_run(witness.get("name", ""))
                else:
                    w_name.add_run(str(witness))

        return self._save_docx_to_bytes(doc)
