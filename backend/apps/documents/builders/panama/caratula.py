"""Panama Caratula (Cover Page) builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_caratula")
class CaratulaBuilder(BaseDocumentBuilder):
    """Build a Panama cover page (Caratula).

    Required context fields:
        entity_name, document_type, date, reference_number.
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "document_type",
        "date",
        "reference_number",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "caratula.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "document_type": context["document_type"],
            "date": context["date"],
            "reference_number": context["reference_number"],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(12)
        style.font.name = "Times New Roman"

        # Spacer
        for _ in range(6):
            doc.add_paragraph()

        # Entity name
        entity_p = doc.add_paragraph()
        entity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = entity_p.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(24)

        doc.add_paragraph()

        # Document type
        type_p = doc.add_paragraph()
        type_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = type_p.add_run(variables["document_type"])
        run.bold = True
        run.font.size = Pt(18)

        doc.add_paragraph()
        doc.add_paragraph()

        # Reference number
        ref_p = doc.add_paragraph()
        ref_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        ref_p.add_run(f"Referencia: {variables['reference_number']}")

        # Date
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        date_p.add_run(f"Fecha: {variables['date']}")

        return self._save_docx_to_bytes(doc)
