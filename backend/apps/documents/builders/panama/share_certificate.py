"""Panama Share Certificate builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_share_certificate")
class ShareCertificateBuilder(BaseDocumentBuilder):
    """Build a Panama share certificate.

    Required context fields:
        entity_name, shareholder_name, share_class, num_shares,
        par_value, certificate_number, issue_date.
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "shareholder_name",
        "share_class",
        "num_shares",
        "certificate_number",
        "issue_date",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "share_certificate.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        return {
            "entity_name": context["entity_name"],
            "shareholder_name": context["shareholder_name"],
            "share_class": context["share_class"],
            "num_shares": context["num_shares"],
            "par_value": context.get("par_value", "N/A"),
            "certificate_number": context["certificate_number"],
            "issue_date": context["issue_date"],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Title
        title = doc.add_heading("CERTIFICADO DE ACCIONES", level=0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER

        doc.add_paragraph()

        # Certificate number
        cert_p = doc.add_paragraph()
        cert_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        run = cert_p.add_run(
            f"Certificado No. {variables['certificate_number']}"
        )
        run.bold = True

        doc.add_paragraph()

        # Entity name
        entity_p = doc.add_paragraph()
        entity_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = entity_p.add_run(variables["entity_name"])
        run.bold = True
        run.font.size = Pt(16)

        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"Se certifica que {variables['shareholder_name']} es titular de "
            f"{variables['num_shares']} acciones de clase "
            f"{variables['share_class']} con un valor nominal de "
            f"{variables['par_value']} cada una."
        )

        doc.add_paragraph(
            f"Fecha de emision: {variables['issue_date']}."
        )

        # Signature lines
        doc.add_paragraph()
        doc.add_paragraph()

        for role in ("Presidente", "Secretario"):
            doc.add_paragraph()
            sig = doc.add_paragraph()
            sig.alignment = WD_ALIGN_PARAGRAPH.CENTER
            sig.add_run("_" * 40)
            label = doc.add_paragraph()
            label.alignment = WD_ALIGN_PARAGRAPH.CENTER
            label.add_run(role)

        return self._save_docx_to_bytes(doc)
