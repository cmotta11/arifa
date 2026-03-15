"""Panama Cover Letter builder."""

from __future__ import annotations

import os

from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt

from ..base import BaseDocumentBuilder, register_builder

_TEMPLATE_DIR = os.path.join(os.path.dirname(__file__), "templates")


@register_builder("panama_cover_letter")
class CoverLetterBuilder(BaseDocumentBuilder):
    """Build a cover letter for a Panama document package.

    Required context fields:
        entity_name, recipient_name, document_list (list), date, sender_name.
    """

    REQUIRED_FIELDS = [
        "entity_name",
        "recipient_name",
        "document_list",
        "date",
        "sender_name",
    ]

    def get_template_path(self) -> str:
        path = os.path.join(_TEMPLATE_DIR, "cover_letter.docx")
        return path if os.path.isfile(path) else ""

    def get_variables(self, context: dict) -> dict:
        document_list = context.get("document_list") or []
        return {
            "entity_name": context["entity_name"],
            "recipient_name": context["recipient_name"],
            "document_list": document_list,
            "date": context["date"],
            "sender_name": context["sender_name"],
        }

    def validate_context(self, context: dict) -> None:
        self._require_fields(context, self.REQUIRED_FIELDS)

    def build(self, context: dict) -> bytes:
        self.validate_context(context)
        variables = self.get_variables(context)

        template_path = self.get_template_path()
        if template_path:
            from docxtpl import DocxTemplate
            import io

            tpl = DocxTemplate(template_path)
            tpl.render(variables)
            buf = io.BytesIO()
            tpl.save(buf)
            buf.seek(0)
            return buf.read()

        doc = Document()
        style = doc.styles["Normal"]
        style.font.size = Pt(11)
        style.font.name = "Times New Roman"

        # Date
        date_p = doc.add_paragraph()
        date_p.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        date_p.add_run(variables["date"])

        doc.add_paragraph()

        # Recipient
        doc.add_paragraph(variables["recipient_name"])

        doc.add_paragraph()

        # Subject
        subject = doc.add_paragraph()
        run = subject.add_run(f"Re: {variables['entity_name']}")
        run.bold = True

        doc.add_paragraph()

        # Body
        doc.add_paragraph(
            f"Estimado/a {variables['recipient_name']}:"
        )
        doc.add_paragraph()
        doc.add_paragraph(
            f"Adjunto encontrara los siguientes documentos correspondientes "
            f"a {variables['entity_name']}:"
        )

        doc.add_paragraph()

        # Document list
        for idx, doc_item in enumerate(variables["document_list"], 1):
            if isinstance(doc_item, dict):
                doc.add_paragraph(
                    f"  {idx}. {doc_item.get('name', doc_item.get('title', 'N/A'))}",
                )
            else:
                doc.add_paragraph(f"  {idx}. {doc_item}")

        doc.add_paragraph()
        doc.add_paragraph(
            "Quedo a su disposicion para cualquier consulta adicional."
        )

        doc.add_paragraph()
        doc.add_paragraph("Atentamente,")

        # Signature
        doc.add_paragraph()
        doc.add_paragraph()

        sig = doc.add_paragraph()
        sig.alignment = WD_ALIGN_PARAGRAPH.LEFT
        sig.add_run("_" * 40)
        name_p = doc.add_paragraph()
        run = name_p.add_run(variables["sender_name"])
        run.bold = True

        return self._save_docx_to_bytes(doc)
